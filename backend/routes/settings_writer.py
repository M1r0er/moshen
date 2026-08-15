"""
墨参 · 设定管理路由
支持多级目录树、AI优化、文档导入、多格式导出。
设定数据存储在 {workspace}/projects/{project_id}/settings/settings_tree.json
"""
import json
import time
import hashlib
import io
import re
from pathlib import Path
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from core.llm_provider import get_llm_provider
from routes.workspace import get_workspace_path, read_text_safe, write_text_safe
from knowledge.project_kb import get_project_kb_manager

router = APIRouter(prefix="/api/settings", tags=["settings"])

# 设定类别
CATEGORIES = {
    "world": "世界观",
    "character": "人物",
    "location": "地区",
    "item": "物品",
    "plot": "剧情",
    "system": "力量体系",
    "other": "其他",
}


# ===== 工具函数 =====

def get_settings_dir(project_id: str) -> Path:
    """获取项目设定目录"""
    ws = get_workspace_path()
    return ws / "projects" / project_id / "settings"


def sanitize_name(name: str) -> str:
    """校验名称，防止路径遍历"""
    if not name or "/" in name or "\\" in name or ".." in name:
        raise HTTPException(400, "无效的名称")
    return name


def ensure_settings_dir(settings_dir: Path) -> Path:
    settings_dir.mkdir(parents=True, exist_ok=True)
    return settings_dir


def get_tree_path(project_id: str) -> Path:
    """获取设定树文件路径"""
    return get_settings_dir(project_id) / "settings_tree.json"


def load_tree(project_id: str) -> dict:
    """加载设定树"""
    project_id = sanitize_name(project_id)
    tree_path = get_tree_path(project_id)
    if tree_path.exists():
        try:
            return json.loads(tree_path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, IOError):
            pass

    # 尝试从旧的 .md 文件迁移
    settings_dir = get_settings_dir(project_id)
    migrated = _migrate_from_flat_files(settings_dir)
    if migrated["nodes"]:
        save_tree(project_id, migrated)
        return migrated

    return {"version": 1, "nodes": []}


def save_tree(project_id: str, tree: dict):
    """保存设定树"""
    project_id = sanitize_name(project_id)
    settings_dir = get_settings_dir(project_id)
    ensure_settings_dir(settings_dir)
    tree_path = get_tree_path(project_id)
    tree_path.write_text(
        json.dumps(tree, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _migrate_from_flat_files(settings_dir: Path) -> dict:
    """从旧的扁平 .md 文件迁移到树结构"""
    nodes = []
    if not settings_dir.exists():
        return {"version": 1, "nodes": nodes}

    cat_map = {"世界观": "world", "人物设定": "character", "剧情大纲": "plot",
               "力量体系": "system", "设定": "other"}
    for f in settings_dir.iterdir():
        if f.is_file() and f.suffix == ".md" and not f.name.startswith("."):
            content = read_text_safe(f)
            name = f.stem
            # 尝试识别类别
            cat = "other"
            for prefix, cat_key in cat_map.items():
                if name.startswith(prefix):
                    cat = cat_key
                    break
            node_id = f"node_{hashlib.md5(name.encode()).hexdigest()[:8]}"
            nodes.append({
                "id": node_id,
                "title": name,
                "category": cat,
                "content": content,
                "children": [],
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            })
    return {"version": 1, "nodes": nodes}


def _gen_node_id() -> str:
    return f"node_{int(time.time()*1000)}_{hashlib.md5(str(time.time()).encode()).hexdigest()[:6]}"


def _find_node(nodes: list, node_id: str) -> dict | None:
    """递归查找节点"""
    for node in nodes:
        if node["id"] == node_id:
            return node
        found = _find_node(node.get("children", []), node_id)
        if found:
            return found
    return None


def _find_parent(nodes: list, node_id: str) -> list | None:
    """递归查找包含目标节点的父列表"""
    for i, node in enumerate(nodes):
        if node["id"] == node_id:
            return nodes
        result = _find_parent(node.get("children", []), node_id)
        if result is not None:
            return result
    return None


def _remove_node(nodes: list, node_id: str) -> dict | None:
    """递归删除并返回被删除的节点"""
    for i, node in enumerate(nodes):
        if node["id"] == node_id:
            return nodes.pop(i)
        removed = _remove_node(node.get("children", []), node_id)
        if removed:
            return removed
    return None


def _count_nodes(nodes: list) -> int:
    """统计节点总数"""
    count = 0
    for node in nodes:
        count += 1
        count += _count_nodes(node.get("children", []))
    return count


def _get_node_path(nodes: list, node_id: str, path: list = None) -> list | None:
    """获取节点的路径（从根到该节点的标题列表）"""
    if path is None:
        path = []
    for node in nodes:
        current_path = path + [node["title"]]
        if node["id"] == node_id:
            return current_path
        result = _get_node_path(node.get("children", []), node_id, current_path)
        if result:
            return result
    return None


def _find_node_by_title(nodes: list, title: str) -> dict | None:
    """按标题递归查找节点（模糊匹配，用于AI设定存写时查找父节点）"""
    for node in nodes:
        if node["title"] == title:
            return node
        found = _find_node_by_title(node.get("children", []), title)
        if found:
            return found
    return None


# 类别中文映射
_CATEGORY_CN_MAP = {
    "世界观": "world", "world": "world",
    "人物": "character", "character": "character",
    "地区": "location", "location": "location",
    "物品": "item", "item": "item",
    "剧情": "plot", "plot": "plot",
    "力量体系": "system", "system": "system",
    "其他": "other", "other": "other",
}


def save_setting_entry(
    project_id: str,
    title: str,
    content: str,
    category: str = "other",
    parent_title: str | None = None,
) -> dict | None:
    """直接保存设定条目（供对话管理器调用，无需 HTTP 请求）

    如果同名设定已存在，自动转为更新而非新建重复条目。

    Args:
        project_id: 项目ID
        title: 设定标题
        content: 设定内容（Markdown）
        category: 设定类别（中英文均可）
        parent_title: 父节点标题（可选，用于创建子设定）

    Returns:
        保存结果 dict，包含 id, title, category 等；失败返回 None
    """
    title = title.strip()
    content = content.strip()
    if not title or not content:
        return None

    cat_key = _CATEGORY_CN_MAP.get(category.strip().lower(), "other")

    try:
        tree = load_tree(project_id)

        # 重名检测：如果同名设定已存在，转为更新
        existing = _find_node_by_title(tree["nodes"], title)
        if existing:
            existing["content"] = content
            existing["category"] = cat_key
            existing["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
            save_tree(project_id, tree)
            return {
                "id": existing["id"],
                "title": title,
                "category": cat_key,
                "category_label": CATEGORIES.get(cat_key, "其他"),
                "updated": True,
            }

        new_node = {
            "id": _gen_node_id(),
            "title": title,
            "category": cat_key,
            "content": content,
            "children": [],
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        }

        if parent_title:
            parent_title = parent_title.strip()
            parent = _find_node_by_title(tree["nodes"], parent_title)
            if parent:
                parent["children"].append(new_node)
            else:
                # 父节点不存在，作为顶级节点
                tree["nodes"].append(new_node)
        else:
            tree["nodes"].append(new_node)

        save_tree(project_id, tree)
        return {
            "id": new_node["id"],
            "title": title,
            "category": cat_key,
            "category_label": CATEGORIES.get(cat_key, "其他"),
        }
    except Exception:
        return None


def get_settings_summary(project_id: str, max_content_len: int = 800) -> str:
    """生成设定树摘要文本，供对话管理器注入上下文

    格式：
      【世界观】
      ▸ 天元大陆
        大陆分为五域，每域有独特的灵气特征...
      ▸ 修炼体系
        ▸ 灵根分类
          天灵根、地灵根、人灵根三级...
    """
    try:
        tree = load_tree(project_id)
        if not tree.get("nodes"):
            return ""

        lines = []
        # 按类别分组
        by_cat = {}
        for node in tree["nodes"]:
            cat = node.get("category", "other")
            by_cat.setdefault(cat, []).append(node)

        def render_node(node, indent):
            title = node.get("title", "")
            content = node.get("content", "").strip()
            prefix = "  " * indent + "▸ "
            lines.append(f"{prefix}{title}")
            if content:
                truncated = content[:max_content_len]
                if len(content) > max_content_len:
                    truncated += "..."
                for cl in truncated.split("\n"):
                    lines.append("  " * (indent + 1) + cl)
            for child in node.get("children", []):
                render_node(child, indent + 1)

        for cat_key in ["world", "character", "location", "item", "plot", "system", "other"]:
            nodes = by_cat.get(cat_key, [])
            if not nodes:
                continue
            cat_label = CATEGORIES.get(cat_key, "其他")
            lines.append(f"【{cat_label}】")
            for node in nodes:
                render_node(node, 0)
            lines.append("")

        return "\n".join(lines).strip()
    except Exception:
        return ""


def update_setting_entry(
    project_id: str,
    title: str,
    content: str | None = None,
    new_title: str | None = None,
    new_category: str | None = None,
) -> dict | None:
    """按标题查找并更新已有设定条目（供对话管理器调用）

    Args:
        project_id: 项目ID
        title: 要查找的设定标题（精确匹配）
        content: 新内容（None 表示不修改内容）
        new_title: 新标题（None 表示不修改标题）
        new_category: 新类别（None 表示不修改类别）

    Returns:
        更新结果 dict；未找到匹配则返回 None
    """
    title = title.strip()
    if not title:
        return None

    try:
        tree = load_tree(project_id)
        node = _find_node_by_title(tree["nodes"], title)
        if not node:
            return None

        if new_title:
            node["title"] = new_title.strip()
        if content is not None:
            node["content"] = content.strip()
        if new_category:
            cat_key = _CATEGORY_CN_MAP.get(new_category.strip().lower(), "other")
            node["category"] = cat_key
        node["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")

        save_tree(project_id, tree)
        return {
            "id": node["id"],
            "title": node["title"],
            "category": node["category"],
            "category_label": CATEGORIES.get(node["category"], "其他"),
            "updated": True,
        }
    except Exception:
        return None


# ===== 请求模型 =====

class OptimizeRequest(BaseModel):
    project_id: str
    content: str
    category: str = "other"


class CreateNodeRequest(BaseModel):
    title: str
    category: str = "other"
    content: str = ""
    parent_id: str | None = None  # None = 顶级节点


class UpdateNodeRequest(BaseModel):
    title: str | None = None
    content: str | None = None
    category: str | None = None


class MoveNodeRequest(BaseModel):
    node_id: str
    new_parent_id: str | None = None  # None = 移到顶级


class ImportFromFileRequest(BaseModel):
    filename: str
    category: str = "other"
    parent_id: str | None = None


class ExportRequest(BaseModel):
    format: str = "txt"  # txt / docx / html / md


# ===== 路由 =====

@router.post("/optimize")
async def optimize_setting(req: OptimizeRequest):
    """AI 优化设定文本"""
    if not req.content.strip():
        raise HTTPException(400, "设定内容不能为空")

    valid_categories = set(CATEGORIES.keys())
    category = req.category if req.category in valid_categories else "other"
    cat_label = CATEGORIES[category]

    llm = get_llm_provider()

    category_prompts = {
        "world": "这是世界观设定。请确保：地理/历史/种族/势力等要素层次分明，设定术语统一，避免冗余描述。突出独特性，让读者快速建立世界认知。",
        "character": "这是人物设定。请确保：性格特征鲜明，外貌描写精炼有力，背景动机逻辑自洽。避免面面俱到，突出核心辨识度。",
        "location": "这是地区设定。请确保：地理特征清晰，文化氛围鲜明，与主线的关联点明确。避免流水账式的罗列。",
        "item": "这是物品设定。请确保：外观功能描述精炼，来历与影响逻辑自洽，与角色/剧情的关联清晰。",
        "plot": "这是剧情大纲。请确保：因果链清晰，冲突递进合理，伏笔埋设与回收有迹可循。用简洁的语言勾勒主线骨架。",
        "system": "这是力量体系设定。请确保：等级划分明确，规则边界清晰，成长路径合理。避免数值堆砌，突出体系特色。",
        "other": "请确保设定文本结构清晰、表达精准。",
    }
    category_hint = category_prompts.get(category, category_prompts["other"])

    prompt = (
        f"你是墨参，一位资深网文设定编辑。请优化以下设定文本。\n\n"
        f"## 优化要求\n"
        f"1. 言简意赅：删除冗余修饰，保留核心信息\n"
        f"2. 逻辑清晰：梳理因果关系，确保设定自洽\n"
        f"3. 表达精准：用准确的设定术语替代口语化描述\n"
        f"4. 保持原意：不改变用户的核心创意和设定方向\n\n"
        f"## 类别指导（{cat_label}）\n{category_hint}\n\n"
        f"## 用户原文\n{req.content}\n\n"
        f"请直接输出优化后的设定文本，不要加任何解释或前言。"
    )
    try:
        result = await llm.generate(
            [{"role": "user", "content": prompt}],
            role="TEXT_MASTER",
            max_tokens=4096,
        )
    except Exception as e:
        raise HTTPException(500, f"LLM 优化失败: {e}")

    return {
        "optimized": result,
        "category": category,
        "original_length": len(req.content),
        "optimized_length": len(result),
    }


@router.get("/{project_id}/tree")
async def get_settings_tree(project_id: str):
    """获取设定树"""
    tree = load_tree(project_id)
    total = _count_nodes(tree["nodes"])
    return {"tree": tree, "total": total}


@router.post("/{project_id}/node")
async def create_node(project_id: str, req: CreateNodeRequest):
    """创建设定节点"""
    if not req.title.strip():
        raise HTTPException(400, "标题不能为空")

    project_id = sanitize_name(project_id)
    tree = load_tree(project_id)

    new_node = {
        "id": _gen_node_id(),
        "title": req.title.strip(),
        "category": req.category if req.category in CATEGORIES else "other",
        "content": req.content,
        "children": [],
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }

    if req.parent_id:
        parent = _find_node(tree["nodes"], req.parent_id)
        if not parent:
            raise HTTPException(404, "父节点不存在")
        parent["children"].append(new_node)
    else:
        tree["nodes"].append(new_node)

    save_tree(project_id, tree)
    return {"success": True, "node": new_node}


@router.put("/{project_id}/node/{node_id}")
async def update_node(project_id: str, node_id: str, req: UpdateNodeRequest):
    """更新设定节点"""
    project_id = sanitize_name(project_id)
    tree = load_tree(project_id)

    node = _find_node(tree["nodes"], node_id)
    if not node:
        raise HTTPException(404, "节点不存在")

    if req.title is not None:
        node["title"] = req.title.strip()
    if req.content is not None:
        node["content"] = req.content
    if req.category is not None:
        node["category"] = req.category if req.category in CATEGORIES else "other"
    node["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")

    save_tree(project_id, tree)
    return {"success": True, "node": node}


@router.delete("/{project_id}/node/{node_id}")
async def delete_node(project_id: str, node_id: str):
    """删除设定节点（及其子节点）"""
    project_id = sanitize_name(project_id)
    tree = load_tree(project_id)

    removed = _remove_node(tree["nodes"], node_id)
    if not removed:
        raise HTTPException(404, "节点不存在")

    save_tree(project_id, tree)
    return {"success": True}


@router.post("/{project_id}/move")
async def move_node(project_id: str, req: MoveNodeRequest):
    """移动节点到新的父节点下"""
    project_id = sanitize_name(project_id)
    tree = load_tree(project_id)

    # 防止将节点移动到自己的子节点下
    if req.new_parent_id == req.node_id:
        raise HTTPException(400, "不能将节点移动到自身下")

    # 检查目标父节点是否是当前节点的子节点
    node = _find_node(tree["nodes"], req.node_id)
    if node and req.new_parent_id:
        descendant = _find_node(node.get("children", []), req.new_parent_id)
        if descendant:
            raise HTTPException(400, "不能将节点移动到其子节点下")

    # 移除节点
    removed = _remove_node(tree["nodes"], req.node_id)
    if not removed:
        raise HTTPException(404, "节点不存在")

    # 插入到新位置
    if req.new_parent_id:
        parent = _find_node(tree["nodes"], req.new_parent_id)
        if not parent:
            # 父节点不存在，放回原处
            tree["nodes"].append(removed)
            save_tree(project_id, tree)
            raise HTTPException(404, "目标父节点不存在")
        parent["children"].append(removed)
    else:
        tree["nodes"].append(removed)

    save_tree(project_id, tree)
    return {"success": True}


@router.post("/{project_id}/import-from-file")
async def import_from_file(project_id: str, req: ImportFromFileRequest):
    """从工作区文档导入设定（AI 提取）"""
    project_id = sanitize_name(project_id)
    kb = get_project_kb_manager()

    # 读取文件内容
    content = None
    # 尝试从 uploads 读取
    project_dir = kb.get_project_dir(project_id)
    if project_dir:
        filepath = project_dir / "uploads" / req.filename
        if filepath.exists():
            content = kb._read_file_for_summary(filepath)

    if not content:
        raise HTTPException(404, "文件不存在")

    # 截取前 8000 字避免超出上下文
    text_content = content[:8000]

    cat_label = CATEGORIES.get(req.category, "其他")
    llm = get_llm_provider()

    prompt = f"""你是墨参，一位资深网文设定编辑。请从以下文档内容中提取{cat_label}相关的设定信息。

## 文档内容
{text_content}

## 提取要求
1. 从文档中识别出与「{cat_label}」相关的设定要素
2. 将提取的内容整理为结构化的设定文本
3. 保留原文的关键信息和独特设定
4. 如果文档中没有直接相关的设定，请基于文档内容进行合理推断和总结
5. 输出格式为 Markdown，使用标题和列表组织内容

请直接输出提取整理后的设定文本，不要加任何解释或前言。"""

    try:
        result = await llm.generate(
            [{"role": "user", "content": prompt}],
            role="TEXT_MASTER",
            max_tokens=4096,
        )
    except Exception as e:
        raise HTTPException(500, f"AI 提取失败: {e}")

    # 自动创建节点
    tree = load_tree(project_id)
    # 用文件名作为标题
    title = Path(req.filename).stem
    new_node = {
        "id": _gen_node_id(),
        "title": f"{title}（AI提取）",
        "category": req.category if req.category in CATEGORIES else "other",
        "content": result,
        "children": [],
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }

    if req.parent_id:
        parent = _find_node(tree["nodes"], req.parent_id)
        if parent:
            parent["children"].append(new_node)
        else:
            tree["nodes"].append(new_node)
    else:
        tree["nodes"].append(new_node)

    save_tree(project_id, tree)
    return {"success": True, "node": new_node, "extracted_content": result}


# ===== 导出功能 =====

def _flatten_tree(nodes: list, level: int = 1) -> list[dict]:
    """将树展平为带层级的列表"""
    result = []
    for node in nodes:
        result.append({
            "title": node["title"],
            "category": node.get("category", "other"),
            "content": node.get("content", ""),
            "level": level,
        })
        result.extend(_flatten_tree(node.get("children", []), level + 1))
    return result


def _export_txt(tree: dict) -> bytes:
    """导出为 TXT"""
    lines = []
    lines.append("墨参 · 设定导出")
    lines.append(f"导出时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("=" * 60)
    lines.append("")

    flat = _flatten_tree(tree["nodes"])
    # 按类别分组
    by_category = {}
    for item in flat:
        cat = item["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(item)

    for cat_key, cat_label in CATEGORIES.items():
        items = by_category.get(cat_key, [])
        if not items:
            continue
        lines.append(f"【{cat_label}】")
        lines.append("-" * 40)
        for item in items:
            indent = "  " * (item["level"] - 1)
            lines.append(f"{indent}{'>' * item['level']} {item['title']}")
            if item["content"]:
                for content_line in item["content"].split("\n"):
                    lines.append(f"{indent}  {content_line}")
            lines.append("")
        lines.append("")

    text = "\n".join(lines)
    return text.encode("utf-8")


def _export_md(tree: dict) -> bytes:
    """导出为 Markdown"""
    lines = []
    lines.append("# 设定导出")
    lines.append(f"> 导出时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")

    flat = _flatten_tree(tree["nodes"])
    by_category = {}
    for item in flat:
        cat = item["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(item)

    for cat_key, cat_label in CATEGORIES.items():
        items = by_category.get(cat_key, [])
        if not items:
            continue
        lines.append(f"## {cat_label}")
        lines.append("")
        for item in items:
            heading_level = min(item["level"] + 2, 6)  # Markdown 最多 6 级标题
            lines.append(f"{'#' * heading_level} {item['title']}")
            if item["content"]:
                lines.append("")
                lines.append(item["content"])
            lines.append("")

    text = "\n".join(lines)
    return text.encode("utf-8")


def _export_html(tree: dict) -> bytes:
    """导出为 HTML"""
    flat = _flatten_tree(tree["nodes"])
    by_category = {}
    for item in flat:
        cat = item["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(item)

    html_parts = [
        "<!DOCTYPE html>",
        '<html lang="zh-CN">',
        "<head>",
        '<meta charset="UTF-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1.0">',
        "<title>设定导出</title>",
        "<style>",
        "body { font-family: 'Microsoft YaHei', sans-serif; max-width: 900px; "
        "margin: 40px auto; padding: 20px; color: #333; line-height: 1.8; }",
        "h1 { color: #7c5cfc; border-bottom: 2px solid #7c5cfc; padding-bottom: 10px; }",
        "h2 { color: #5b3fd6; margin-top: 30px; border-left: 4px solid #7c5cfc; padding-left: 12px; }",
        "h3, h4, h5, h6 { color: #444; margin-top: 20px; }",
        ".content { background: #f8f9fa; padding: 12px 16px; border-radius: 6px; "
        "margin: 8px 0; white-space: pre-wrap; }",
        ".meta { color: #999; font-size: 12px; }",
        ".category { background: #7c5cfc; color: #fff; padding: 4px 12px; "
        "border-radius: 4px; display: inline-block; margin: 20px 0 10px; }",
        "</style>",
        "</head>",
        "<body>",
        "<h1>设定导出</h1>",
        f'<p class="meta">导出时间：{time.strftime("%Y-%m-%d %H:%M:%S")}</p>',
    ]

    for cat_key, cat_label in CATEGORIES.items():
        items = by_category.get(cat_key, [])
        if not items:
            continue
        html_parts.append(f'<div class="category">{cat_label}</div>')
        for item in items:
            tag = f"h{min(item['level'] + 2, 6)}"
            html_parts.append(f"<{tag}>{item['title']}</{tag}>")
            if item["content"]:
                # 简单转义 HTML
                escaped = item["content"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_parts.append(f'<div class="content">{escaped}</div>')

    html_parts.append("</body>")
    html_parts.append("</html>")
    return "\n".join(html_parts).encode("utf-8")


def _export_docx(tree: dict) -> bytes:
    """导出为 Word 文档"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()

    # 标题
    title_para = doc.add_heading("设定导出", level=0)
    title_para.alignment = 1  # 居中

    # 导出时间
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(f"导出时间：{time.strftime('%Y-%m-%d %H:%M:%S')}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    flat = _flatten_tree(tree["nodes"])
    by_category = {}
    for item in flat:
        cat = item["category"]
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(item)

    for cat_key, cat_label in CATEGORIES.items():
        items = by_category.get(cat_key, [])
        if not items:
            continue

        doc.add_heading(cat_label, level=1)

        for item in items:
            # Word 标题级别（1-9），限制在 2-6
            heading_level = min(item["level"] + 1, 6)
            doc.add_heading(item["title"], level=heading_level)

            if item["content"]:
                # 将内容按段落添加
                for para_text in item["content"].split("\n"):
                    if para_text.strip():
                        doc.add_paragraph(para_text)

                doc.add_paragraph("")  # 空行分隔

    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@router.get("/{project_id}/export")
async def export_settings(project_id: str, format: str = "txt"):
    """导出设定为指定格式"""
    project_id = sanitize_name(project_id)
    tree = load_tree(project_id)

    if not tree["nodes"]:
        raise HTTPException(400, "没有设定内容可导出")

    fmt = format.lower()
    timestamp = time.strftime("%Y%m%d_%H%M%S")

    if fmt == "txt":
        content = _export_txt(tree)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename=settings_{timestamp}.txt"},
        )
    elif fmt == "md":
        content = _export_md(tree)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=settings_{timestamp}.md"},
        )
    elif fmt == "html":
        content = _export_html(tree)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="text/html",
            headers={"Content-Disposition": f"attachment; filename=settings_{timestamp}.html"},
        )
    elif fmt == "docx":
        content = _export_docx(tree)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=settings_{timestamp}.docx"},
        )
    else:
        raise HTTPException(400, f"不支持的格式: {fmt}")


# ===== 兼容旧接口 =====

@router.get("/{project_id}")
async def list_settings_legacy(project_id: str):
    """兼容旧接口：从树结构生成文件列表"""
    tree = load_tree(project_id)
    files = []
    flat = _flatten_tree(tree["nodes"])
    for item in flat:
        files.append({
            "filename": f"{item['title']}.md",
            "size": len(item["content"]),
            "modified": time.strftime("%Y-%m-%d %H:%M:%S"),
        })
    return {"files": files}


@router.post("/{project_id}/legacy-save")
async def save_setting_legacy(project_id: str, req: dict):
    """兼容旧接口：保存为树节点"""
    project_id = sanitize_name(project_id)
    filename = req.get("filename", "").replace(".md", "")
    content = req.get("content", "")
    if not filename:
        raise HTTPException(400, "文件名不能为空")

    tree = load_tree(project_id)
    new_node = {
        "id": _gen_node_id(),
        "title": filename,
        "category": "other",
        "content": content,
        "children": [],
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    tree["nodes"].append(new_node)
    save_tree(project_id, tree)
    return {"success": True, "filename": f"{filename}.md"}
